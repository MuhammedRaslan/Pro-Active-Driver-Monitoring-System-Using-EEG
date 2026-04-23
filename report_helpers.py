"""
Report Helper Functions — Styles, utilities, and constants for report generation.
"""
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# ── Paths ────────────────────────────────────────────────────────────────────
BASE = os.path.dirname(os.path.abspath(__file__))
FIG_DIR = os.path.join(BASE, "Files")
ART_DIR = r"C:\Users\muham\.gemini\antigravity\brain\1c14171f-ca04-48e0-b3ec-e3ec885ecd78"
OUT_PATH = os.path.join(BASE, "Capstone_Report.docx")

GENERATED_FIGS = {
    "sys_arch": os.path.join(ART_DIR, [f for f in os.listdir(ART_DIR) if f.startswith("system_architecture")][0]) if os.path.isdir(ART_DIR) else "",
    "pred_algo": os.path.join(ART_DIR, [f for f in os.listdir(ART_DIR) if f.startswith("prediction_algorithm")][0]) if os.path.isdir(ART_DIR) else "",
    "alert_tiers": os.path.join(ART_DIR, [f for f in os.listdir(ART_DIR) if f.startswith("alert_system")][0]) if os.path.isdir(ART_DIR) else "",
    "dataset": os.path.join(ART_DIR, [f for f in os.listdir(ART_DIR) if f.startswith("dataset_setup")][0]) if os.path.isdir(ART_DIR) else "",
    "headrest": os.path.join(ART_DIR, [f for f in os.listdir(ART_DIR) if f.startswith("headrest_schematic")][0]) if os.path.isdir(ART_DIR) else "",
    "dataflow": os.path.join(ART_DIR, [f for f in os.listdir(ART_DIR) if f.startswith("data_flow")][0]) if os.path.isdir(ART_DIR) else "",
}

EXISTING_FIGS = {
    "sys_arch_existing": os.path.join(FIG_DIR, "System_Architecture..png"),
    "headrest_flow": os.path.join(FIG_DIR, "Headrest_Flow.png"),
    "ml_flow": os.path.join(FIG_DIR, "ML_Flowchart.png"),
    "raw_awake": os.path.join(FIG_DIR, "Raw_Awake.png"),
    "raw_drowsy": os.path.join(FIG_DIR, "Raw_Drowsy.png"),
    "raw_compare": os.path.join(FIG_DIR, "Raw_Comparison.png"),
}


def create_doc():
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    pf = style.paragraph_format
    pf.line_spacing = 1.5
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_after = Pt(6)

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1)

    return doc


def add_chapter(doc, number, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.page_break_before = True
    run = p.add_run(f"CHAPTER {number}")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Times New Roman'
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p2.paragraph_format.space_after = Pt(18)
    run2 = p2.add_run(title.upper())
    run2.bold = True
    run2.font.size = Pt(14)
    run2.font.name = 'Times New Roman'


def add_section(doc, number, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(f"{number}  {title}")
    run.bold = True
    run.font.size = Pt(13)
    run.font.name = 'Times New Roman'


def add_subsection(doc, number, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f"{number}  {title}")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'


def add_para(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    return p


def add_figure(doc, path, caption, fig_num, width=5.5):
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Inches(width))
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"[IMAGE NOT FOUND: {os.path.basename(path)}]")
        run.font.color.rgb = RGBColor(255, 0, 0)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(12)
    run = cap.add_run(f"Figure {fig_num}: {caption}")
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'


def add_placeholder(doc, label, fig_num, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(f"\n[SCREENSHOT PLACEHOLDER: {label}]\n")
    run.font.color.rgb = RGBColor(128, 128, 128)
    run.font.size = Pt(11)
    run.italic = True
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(12)
    run = cap.add_run(f"Figure {fig_num}: {caption}")
    run.bold = True
    run.font.size = Pt(10)


def add_table(doc, headers, rows, table_num, caption):
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(12)
    run = cap.add_run(f"Table {table_num}: {caption}")
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
                r.font.name = 'Times New Roman'

    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.size = Pt(10)
                    r.font.name = 'Times New Roman'

    doc.add_paragraph()  # spacer
    return table


def add_code_block(doc, code, lang="python"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)


def page_break(doc):
    doc.add_page_break()
